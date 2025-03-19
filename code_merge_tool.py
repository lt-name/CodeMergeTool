# MIT License
#
# Copyright (c) 2024 Jiacheng Ni
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.


import os
import json
import re
from docx import Document

# 加载配置文件，忽略 "_comment" 字段
with open('config_file.json', 'r', encoding='utf-8') as f:
    config = json.load(f)

# 去除 "_comment" 字段，仅保留实际配置项
config = {key: value for key, value in config.items() if not key.startswith("_comment")}

# 提取配置项
project_root = config["PROJECT_ROOT"]                 # 项目根目录
output_dir = config["OUTPUT_DIR"]                     # 输出目录
all_code_file = os.path.join(output_dir, config["ALL_CODE_FILE"])
output_first_doc = os.path.join(output_dir, config["OUTPUT_FIRST_DOC"])
output_last_doc = os.path.join(output_dir, config["OUTPUT_LAST_DOC"])
file_types = config["FILE_TYPES"]

# 创建输出目录（如果不存在）
os.makedirs(output_dir, exist_ok=True)

# 正则表达式匹配注释和空行
single_line_start= r'//.*'                                # 匹配行后注释开始
single_line_comment = re.compile(r'^\s*//')               # 匹配单行注释
multi_line_comment_start = re.compile(r'^\s*/\*')         # 匹配多行注释的开始
multi_line_comment_end = re.compile(r'.*\*/\s*$')         # 匹配多行注释的结束

# 创建或清空 All_code 文件
with open(all_code_file, 'w', encoding='utf-8') as outfile:
    # 遍历项目根目录中的所有文件
    for foldername, subfolders, filenames in os.walk(project_root):
        for filename in filenames:
            # 获取每个文件的完整路径
            file_path = os.path.join(foldername, filename)

            # 只处理指定的文件类型
            if any(filename.endswith(ext) for ext in file_types):
                try:
                    with open(file_path, 'r', encoding='utf-8') as infile:
                        in_multiline_comment = False  # 标识是否在多行注释块中

                        for line in infile:
                            # 跳过空行
                            if line.strip() == '':
                                continue

                            # 检查是否进入或退出多行注释
                            if in_multiline_comment:
                                if multi_line_comment_end.search(line):
                                    in_multiline_comment = False
                                continue
                            elif multi_line_comment_start.match(line):
                                in_multiline_comment = True
                                continue

                            # 跳过单行注释
                            if single_line_comment.match(line):
                                continue

                            # 去除行后注释
                            line = re.sub(single_line_start, '', line)

                            # 将非注释、非空行写入 All_code 文件
                            outfile.write(line)
                except UnicodeDecodeError:
                    print(f"⚠️ 无法读取文件（编码错误），跳过该文件：{file_path}")

print(f'所有非注释、非空行代码已成功合并到 {all_code_file}')

# 从 All_code 文件中提取前1500行和最后1500行
cleaned_lines = []
with open(all_code_file, 'r', encoding='utf-8') as f:
    in_multiline_comment = False  # 重置多行注释标志
    for line in f:
        # 过滤空行和注释行
        if line.strip() == '' or single_line_comment.match(line):
            continue

        # 跳过多行注释
        if in_multiline_comment:
            if multi_line_comment_end.search(line):
                in_multiline_comment = False
            continue
        if multi_line_comment_start.match(line):
            in_multiline_comment = True
            continue

        # 添加非注释的行
        cleaned_lines.append(line)

# 获取清理后的总行数
total_cleaned_lines = len(cleaned_lines)
print(f'清理后的代码总行数: {total_cleaned_lines}')

# 提取前1500行和最后1500行
first_1500_lines = cleaned_lines[:1500] if total_cleaned_lines >= 1500 else cleaned_lines
last_1500_lines = cleaned_lines[-1500:] if total_cleaned_lines >= 1500 else cleaned_lines

# 创建两个Word文档，分别保存前1500行和最后1500行
first_doc = Document()
last_doc = Document()

# 将前1500行写入第一个文档，确保不添加空行
for line in first_1500_lines:
    if line.strip():  # 确保不添加空行
        first_doc.add_paragraph(line.strip())

# 将最后1500行写入第二个文档，确保不添加空行
for line in last_1500_lines:
    if line.strip():  # 确保不添加空行
        last_doc.add_paragraph(line.strip())

# 保存两个Word文档
first_doc.save(output_first_doc)
last_doc.save(output_last_doc)

print("前1500行和最后1500行的代码已分别保存到Word文档，并去除空行。")
